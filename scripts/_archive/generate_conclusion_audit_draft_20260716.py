from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "conclusion_2026"
OUT_PATH = OUT_DIR / "结项材料预审与结题报告填写底稿_20260716.docx"


BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "E7E6E6"
LIGHT_RED = "FCE4D6"
LIGHT_YELLOW = "FFF2CC"
LIGHT_GREEN = "E2F0D9"
DARK = RGBColor(31, 41, 55)
MUTED = RGBColor(89, 89, 89)
RED = RGBColor(192, 0, 0)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_col_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    paragraph.add_run(" 页")


def add_field_table(doc: Document, rows: list[tuple[str, str]], widths=(3.4, 12.0)):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_BLUE)
        cells[0].paragraphs[0].runs[0].bold = True
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
        prevent_row_split(table.rows[-1])
    set_col_widths(table, list(widths))
    return table


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, label in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = label
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(cell)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[idx])
        prevent_row_split(table.rows[-1])
    if widths:
        set_col_widths(table, widths)
    return table


def add_bullet(doc: Document, text: str, level: int = 0, color: RGBColor | None = None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    if color:
        run.font.color.rgb = color
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def add_note(doc: Document, text: str, kind: str = "warn"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    fill = LIGHT_RED if kind == "danger" else LIGHT_YELLOW if kind == "warn" else LIGHT_GREEN
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=140, end=140)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    if kind == "danger":
        run.font.color.rgb = RED
    return table


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(5)

    for name, size, font, color in (
        ("Title", 20, "方正小标宋简体", BLUE),
        ("Heading 1", 15, "黑体", BLUE),
        ("Heading 2", 12.5, "黑体", DARK),
        ("Heading 3", 11, "黑体", DARK),
    ):
        style = styles[name]
        style.font.name = font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color) if isinstance(color, str) else color
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        styles[style_name].font.name = "宋体"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        styles[style_name].font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("五邑大学实验室安全课题结项材料 · 填写底稿")
    hr.font.size = Pt(9)
    hr.font.color.rgb = MUTED
    add_page_number(section.footer.paragraphs[0])


def build() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_document_defaults(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(45)
    r = p.add_run("结项材料预审与结题报告填写底稿")
    r.bold = True
    r.font.name = "方正小标宋简体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "方正小标宋简体")
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于 Dify 搭建 RAG 增强的大语言模型实验室安全小助手系统")
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph()
    cover = add_field_table(
        doc,
        [
            ("课题负责人", "龙华秋"),
            ("所在单位", "电子与信息工程学院"),
            ("课题类别", "一般课题（校级）"),
            ("立项申报日期", "2025年6月25日"),
            ("原预计完成日期", "2026年6月30日"),
            ("底稿核验日期", "2026年7月16日"),
            ("学校结项截止", "2026年8月31日；发送至 sbc@wyu.edu.cn"),
        ],
    )
    cover.rows[0].cells[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    add_note(
        doc,
        "文档状态：事实核验底稿，非学校官方附件2。官方《项目结题报告（简版）》及项目编号尚未取得；在迁入官方模板、核实经费、补齐签字盖章前不得直接提交。",
        "danger",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("依据：学校结项通知、原立项书、当前仓库代码与 2026-07-16 本机实测")
    r.font.size = Pt(9.5)
    r.font.color.rgb = MUTED

    doc.add_page_break()

    doc.add_heading("一、结项通知要求与预审结论", level=1)
    add_field_table(
        doc,
        [
            ("通知", "《关于开展2025年度校内实践类项目结题验收及2026年度项目申报工作的通知》，2026年7月13日发布"),
            ("结项对象", "2025年以来已立项且项目周期已结束的项目"),
            ("提交材料", "附件2《项目结题报告（简版）》"),
            ("截止时间", "2026年8月31日"),
            ("提交方式", "电子材料发送至 sbc@wyu.edu.cn"),
            ("咨询", "实验室与设备管理处覃老师，电话 3296175"),
            ("原通知链接", "https://www.wyu.edu.cn/shebeichu/info/1003/3373.htm"),
        ],
    )

    doc.add_heading("预审总判断", level=2)
    add_note(
        doc,
        "项目已经形成可演示原型和完整代码/数据骨架，具备“结项整改后提交”的基础；但目前不宜原样宣称所有量化指标已完成。最高风险在于 Dify 主链路离线、连续试运行证据中断、降级问答有效率不足、知识条目绝大多数仍未标记为 reviewed。",
        "warn",
    )
    add_matrix(
        doc,
        ["结论维度", "判断", "说明"],
        [
            ["可演示性", "基本具备", "React 页面、FastAPI 后端、本地检索、规则拦截均可运行；高风险问题“离心机运转时能否开盖”被 R-023 拒绝。"],
            ["官方表格", "待取得", "学校附件下载触发验证码；项目清单中的项目编号也待取得。"],
            ["核心数量", "数量达成", "知识库 CSV 为 3009 行、29 列，source_url/source_title 均无空值。"],
            ["知识质量", "未充分举证", "仅 38 条状态为 reviewed；1871 条 draft、1100 条 external_draft。"],
            ["有效回答率", "当前不达标", "Dify 离线降级评测中，保守人工审计约 22/50 基本可用；不可作为 >95% 的完成证明。"],
            ["响应性能", "仅本地链路达标", "50题本地降级链路平均 138ms；Dify SSE 首字节和完整响应当前未复测。"],
            ["7×24/三个月", "未完成", "现有 26 个快照截至 2026-06-15；监控任务路径失效，未形成连续三个月证据。"],
            ["最终建议", "整改后据实结项", "把“完成原型、数量达标、规则拦截可演示”作为主成果；将未完成指标列为局限和后续计划。"],
        ],
        [2.8, 2.5, 10.1],
    )

    doc.add_heading("二、官方结题表通用字段填写底稿", level=1)
    add_note(doc, "以下内容用于迁入学校附件2；带【待核实】的字段必须由负责人确认后再提交。", "info")

    doc.add_heading("2.1 基本信息", level=2)
    add_field_table(
        doc,
        [
            ("项目编号", "【待从学校附件1项目清单取得】"),
            ("项目名称", "基于Dify搭建RAG增强的大语言模型实验室安全小助手系统"),
            ("项目类别", "一般课题"),
            ("负责人", "龙华秋"),
            ("负责人单位", "电子与信息工程学院"),
            ("项目周期", "2025年7月—2026年6月（以立项批文为准）"),
            ("预期成果", "软件系统、技术路线"),
            ("批准经费", "5000元（以财务/立项批文为准）"),
        ],
    )

    doc.add_heading("2.2 主要参加人员", level=2)
    add_matrix(
        doc,
        ["姓名", "原立项分工", "结项材料建议表述"],
        [
            ["谢惠敏", "安全问答设计", "安全问答场景设计与审核协助"],
            ["伍良恒", "平台搭建", "Dify/部署环境搭建"],
            ["伍栢深", "大模型研究", "大模型与 RAG 方案研究"],
            ["李烨", "安全问答指导", "实验室安全业务指导"],
            ["洪智勇", "顶层设计", "项目总体方案与数据分析指导"],
            ["冼健邦", "智能体开发", "智能体流程开发"],
            ["陈浩林", "前端设计", "React 演示界面设计"],
            ["唐泽红", "文档设计", "项目文档与材料整理"],
            ["邓胜杰", "后端开发", "FastAPI 后端与接口开发"],
            ["张丽虹", "资料收集", "安全资料与知识来源整理"],
        ],
        [3.0, 4.0, 8.4],
    )
    p = doc.add_paragraph("说明：最终分工必须由负责人根据真实参与情况确认，不以代码仓库推断替代成员本人确认。")
    p.runs[0].font.color.rgb = RED

    doc.add_heading("2.3 项目完成情况（可直接粘贴的建议稿）", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run(
        "本项目围绕高校实验室安全问答场景，完成了基于 FastAPI、React、Dify 与检索增强生成（RAG）的实验室安全小助手原型。系统实现了用户提问、本地知识库检索、安全规则匹配、Dify 上游调用、结构化回答、引用展示和降级处理等主要功能；建设了 3009 条、29 字段的结构化知识片段库，并建立 safety_rules.yaml 单一规则文件和 25 条安全规则。项目已具备本地部署、功能演示、批量评测和运行监测的基本能力。"
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run(
        "经 2026 年 7 月 16 日复核，知识库数量和软件原型已形成，但部分原定量化目标尚未形成充分的验收证据：Dify 主链路当日未联通，降级链路的问答相关性仍需优化；知识条目中仅 38 条标记为 reviewed；连续 7×24 三个月试运行监测尚未完成。因此，本报告据实将项目定位为“核心原型与技术路线完成、知识数量达标、质量审查和长期试运行仍需补强”，不对整体 99% 准确率或连续三个月运行作无证据承诺。"
    )

    doc.add_heading("2.4 主要研究工作与技术路线", level=2)
    for text in (
        "知识工程：统一 29 列知识库 Schema，形成 3009 条结构化片段，保存来源组织、标题、版本、日期和 URL 等可追溯字段。",
        "系统架构：构建 React 单页应用与 FastAPI 服务，提供健康检查、知识检索、问答、反馈、Dify 参数代理和 Dify Chat SSE 代理。",
        "安全控制：使用 host 白名单约束 Dify API 地址；服务器端注入 Dify App Key，不透传入站 Authorization；对危险操作和应急场景进行规则拦截或结构化答复。",
        "降级策略：当 Dify 不可用时，尝试本地快速检索、规则引擎和结构化兜底；当前已验证可返回响应，但相关性仍需整改。",
        "评测与运维：提供 50 题批量评测、质量门禁、运行快照和周报脚本；本次复核发现监控调度路径漂移，需要重新注册。",
    ):
        add_bullet(doc, text)

    doc.add_heading("2.5 原定指标与当前事实对照", level=2)
    add_matrix(
        doc,
        ["原定指标", "当前可证明事实", "结项填写建议"],
        [
            ["知识库≥3000条", "CSV 3009 行，29 列；725 个唯一 source_url，1284 个唯一 source_title；URL/标题均非空。", "写“3000+结构化知识片段，来自400+权威来源”，不要写“3009份独立文档”。"],
            ["有效回答率>95%", "已有材料口径相互矛盾；本次 Dify 离线降级审计仅约 44% 基本可用。", "暂不宣称完成；恢复 Dify 后重新评测并由专家签字。"],
            ["专业准确率>99%", "已有 L1/L2/L3 自评记录，但专家 L4 尚待完成；当前 50题存在明显错配。", "写“已建立分层评测机制”，不要写“整体准确率99%”。"],
            ["响应<3秒", "本地降级 50题平均138ms、最大244ms；Dify 在线 SSE 首字节未在本次复核中测得。", "限定为“本地快速通道<300ms”；在线主链路待复测。"],
            ["7×24服务/试运行3个月", "仅26个旧快照，健康和Dify可达率均30.8%，最新快照截至2026-06-15。", "如实写“已建立监测机制，连续三个月证据未完成”；不得倒填。"],
            ["软件系统、技术路线", "FastAPI+React原型、Dify代理、RAG检索、规则引擎和相关文档均存在。", "可写“已完成核心原型和技术路线”。"],
        ],
        [3.0, 7.1, 5.3],
    )

    doc.add_heading("2.6 主要成果与创新点（建议稿）", level=2)
    for text in (
        "形成可本地演示的实验室安全问答软件原型，覆盖问答、检索、引用、安全拦截、反馈与知识库态势展示。",
        "形成 29 字段结构化知识库 Schema 和 3009 条知识片段，保留来源可追溯字段。",
        "将规则引擎置于生成式问答之前，对危险操作进行确定性阻断，并为应急场景提供结构化处置模板。",
        "建立 Dify 不可用时的本地降级链路，保证接口可响应；本次审计同时识别出降级相关性不足，已形成明确整改方向。",
        "形成质量门禁、批量评测和运行监测工具，为后续专家评审和持续优化提供基础。",
    ):
        add_bullet(doc, text)
    p = doc.add_paragraph("不建议继续使用的未经证明表述：首创、生产级、200并发500ms、万级注入阻断99.2%、知识自动同步1小时、Dify企业授权、校园SSO已接入、签署试运行协议、LoRA联邦协同已完成。")
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = RED

    doc.add_heading("2.7 经费决算填写底稿", level=2)
    add_matrix(
        doc,
        ["经费科目", "原计划金额（元）", "实际支出（元）", "票据/说明"],
        [
            ["Dify/模型 API 调用", "2000", "【待财务核实】", "【待补发票或后台账单】"],
            ["云服务器/部署", "1500", "【待财务核实】", "【待补合同、发票或校内资源说明】"],
            ["知识库整理与标注", "1000", "【待财务核实】", "【不得虚构劳务支出】"],
            ["文档与成果材料", "500", "【待财务核实】", "【待补打印/认证票据】"],
            ["合计", "5000", "【待财务核实】", "实际支出、结余和退回金额以财务为准"],
        ],
        [4.5, 3.3, 3.3, 4.3],
    )
    add_note(doc, "经费数据是目前无法从代码仓库核实的硬性字段。未取得财务流水前，不能把计划预算当成实际决算。", "danger")

    doc.add_heading("2.8 存在问题及后续计划（可直接粘贴的建议稿）", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run(
        "项目当前主要不足为知识条目专家审核覆盖不足、Dify 在线链路稳定性和连续监测证据不足、离线检索在部分专业设备与化学品问题上存在错配。后续将按高风险优先原则完成知识条目复核，针对错配问题增加问题—答案回归集和检索阈值/重排策略，恢复 Dify 服务后重新开展 50 题及扩展题集评测，并由实验室安全专家独立评分；同时修复运行监控任务路径，持续积累真实运行数据，不追溯补填。"
    )

    doc.add_heading("2.9 审核意见与签章", level=2)
    add_field_table(
        doc,
        [
            ("项目负责人确认", "签字：________________    日期：____年__月__日"),
            ("所在单位意见", "负责人签字：____________    单位盖章：____________    日期：____年__月__日"),
            ("主管部门意见", "【按学校附件2原表保留，不代填】"),
        ],
    )

    doc.add_page_break()
    doc.add_heading("三、2026-07-16 代码与端到端实测记录", level=1)
    add_matrix(
        doc,
        ["检查项", "结果", "证据/影响"],
        [
            ["Git工作树", "原始状态干净；master较origin超前2个提交", "本次检查新增结项评测与底稿文件，未修改业务源码。"],
            ["Python质量门禁", "通过", "使用项目实际解释器运行 scripts/quality_gate.py，通过29列表头等检查。"],
            ["Python语法扫描", "1个失败", "scripts/generate_knowledge_bulk_v2.py 含大量弯引号，约第60行触发语法错误。"],
            ["前端构建", "通过", "临时目录执行 tsc -b && vite build 成功，651个模块完成构建。"],
            ["前端Lint", "失败", "package.json定义eslint命令，但devDependencies未安装eslint及配置依赖。"],
            ["本地后端", "通过", "127.0.0.1:8091 /health、/api/meta、/api/search、/api/chat 可响应。"],
            ["高风险规则", "通过", "“离心机运转时能不能打开盖子？”被R-023规则拦截，decision=rule_blocked。"],
            ["50题传输稳定性", "50/50 HTTP成功", "平均138ms，中位139ms，最大244ms；此结果仅说明降级链路可返回，不代表答案正确。"],
            ["50题人工可用性", "约22/50基本可用", "28题为泛化模板或明显错配；本次为保守审计，不是正式专家评分。"],
            ["Dify主链路", "未通过", "Docker/Dify当前不可达；未获得启动Docker授权，未执行恢复。"],
            ["来源链接抽检", "9/20直接可达", "其余主要因CDC/OSHA TLS或代理连接失败；不能全部判定为死链，但现场抽查风险较高。"],
            ["运行监测", "未通过", "计划任务仍指向旧目录D:\\newwork\\lab-safe-assistant-dify-rag，当前仓库已迁至D:\\newwork\\Security\\...。"],
        ],
        [3.2, 3.3, 8.9],
    )

    doc.add_heading("明显 Bug / 风险清单", level=2)
    add_matrix(
        doc,
        ["优先级", "问题", "影响", "建议"],
        [
            ["P0", "Dify不可达，降级答案大量泛化或错配", "主链路E2E和>95%有效率无法证明；安全问答可能误导。", "恢复Dify；对50题逐项复评；低置信度时拒答/转人工，不要强行本地直出。"],
            ["P0", "7×24监控任务路径失效", "自2026-06-15后无有效快照；三个月证据链中断。", "修正并重新注册计划任务，先做单次验证，再按真实时间持续采集。"],
            ["P0", "2971/3009条仍为draft状态", "不能支撑“全部已审核”或整体99%准确率。", "优先审核高风险/高频条目，保留reviewer、日期和抽样记录。"],
            ["P1", "演示登录未配置密码时默认放行", "任意字符串可登录；页面仍显示“密码认证已启用/已加密”。", "提交演示前强制配置DEMO_PASSWORD；未配置时服务启动失败或UI明确标注未启用。"],
            ["P1", "/health始终返回ok，即使Dify不可达", "前端显示“服务在线”，掩盖主链路降级。", "返回明确degraded状态，并在顶部同时显示“本地在线/Dify离线”。"],
            ["P1", "meta_routes的requests.get未关闭Response", "反复健康检查可能造成连接资源泄漏。", "使用with requests.get(...)或finally close。"],
            ["P1", "generate_knowledge_bulk_v2.py语法损坏", "活跃脚本无法运行，影响复现和数据维护。", "统一替换弯引号并补最小回归测试。"],
            ["P2", "前端lint脚本不可执行", "缺少静态质量门禁。", "补eslint依赖/配置，或删除失效脚本并采用当前工具链。"],
            ["P2", "8088/8091端口和Python依赖漂移", "README、评测脚本、运行任务口径不一致；当前环境版本高于requirements。", "统一由单一环境变量读取端口；建立虚拟环境并锁定安装。"],
            ["P2", "无正式pytest回归套件", "规则、SSRF、资源关闭和降级策略易回归。", "至少覆盖认证、健康降级、规则拦截、Dify代理、fallback和Schema。"],
        ],
        [1.4, 4.3, 5.0, 4.7],
    )

    doc.add_heading("典型错配样例", level=2)
    add_matrix(
        doc,
        ["问题", "当前回答问题", "判定"],
        [
            ["钠金属如何安全储存和处置", "返回丙酮的闪点、通风和废液信息。", "明显错误"],
            ["HPLC安全注意事项", "返回GC-MS载气、真空和离子源操作。", "明显错误"],
            ["生物废弃物如何处置", "返回恒温培养箱使用方法。", "明显错误"],
            ["高压电源如何上电", "返回高压灭菌锅托盘加水与干热循环。", "明显错误"],
            ["液氮安全事项", "主要返回设备退役与联系EHS，未完整回答冻伤、窒息和密闭增压。", "不可用"],
            ["化学中毒急救/触电急救/火灾处置", "多题复用“暂停—核对PPE—按SOP”通用模板，缺少关键急救/撤离步骤。", "高风险不可用"],
        ],
        [5.0, 7.8, 2.6],
    )

    doc.add_page_break()
    doc.add_heading("四、提交前必做清单", level=1)
    for text in (
        "取得学校附件1项目清单，确认项目编号、正式项目名称、立项日期和批准经费。",
        "取得学校附件2《项目结题报告（简版）》原版，把本底稿迁入，不改变官方表格结构。",
        "负责人核实参与人员、实际分工、经费实际支出、票据和结余；补齐签字盖章。",
        "经授权启动Docker/Dify，验证8081端口和/v1/chat-messages SSE主链路。",
        "修复或临时禁用错误本地直出：遇低置信度应明确拒答并转人工，不应输出错配SOP。",
        "完成50题人工复评；高风险错误必须全部清零，再由至少1名实验室安全专家复核并签名。",
        "重新注册运行监测任务并验证新路径；报告中只使用真实采集日期，不倒填三个月数据。",
        "把知识审核状态、reviewer和日期补齐；优先覆盖应急、危化品、用电和高风险设备。",
        "复核20条答辩抽查用source_url，准备可离线展示的来源标题、机构、日期和关键页截图。",
        "打包软件、技术路线、操作说明、评测报告、运行截图、知识库统计和源码版本号作为附件。",
        "将最终材料发送至sbc@wyu.edu.cn前，由负责人进行一次逐项事实签字确认。",
    ):
        add_number(doc, text)

    doc.add_heading("建议附件目录", level=2)
    add_matrix(
        doc,
        ["序号", "附件", "当前状态"],
        [
            ["1", "学校官方项目结题报告（简版）", "待下载并填写"],
            ["2", "系统操作说明/README（修订版）", "已有但需统一端口与完成度口径"],
            ["3", "技术路线图与系统架构说明", "已有材料可整理"],
            ["4", "3009条知识库统计与Schema质量门禁报告", "可生成；需补审核覆盖说明"],
            ["5", "50题正式评测与专家评分表", "传输评测已生成；专家评分待完成"],
            ["6", "规则引擎清单与高风险拦截演示截图", "25条规则；可补截图"],
            ["7", "连续试运行报告", "证据不足；只能提交现有真实快照与情况说明"],
            ["8", "经费决算与票据汇总", "待负责人/财务提供"],
            ["9", "源码版本说明与部署包", "仓库已有；提交前打标签/导出"],
        ],
        [1.4, 9.0, 5.0],
    )

    doc.add_heading("五、待负责人确认事项", level=1)
    add_note(doc, "以下信息无法从学校网页、立项书和代码仓库独立核实，必须由负责人补充。", "warn")
    for text in (
        "是否同意处理学校下载验证码，以取得附件1和附件2？",
        "是否同意启动Docker Desktop并恢复本地Dify 8081进行在线端到端复测？",
        "正式项目编号、批准经费、立项批文日期是否与申报书一致？",
        "5000元经费的实际支出、票据、结余分别是多少？",
        "是否存在真实的试运行起止时间、用户单位、试用记录或反馈表？",
        "是否已取得论文、软件著作权、获奖、应用证明等新增成果？如无，不应临时补写。",
        "由哪位实验室安全专家承担最终50题评分并签字？",
    ):
        add_bullet(doc, text)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(15)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("—— 本底稿结束 ——")
    r.bold = True
    r.font.color.rgb = MUTED

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build())
